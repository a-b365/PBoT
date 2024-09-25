import io
import fitz
import docx
import gridfs
import urllib
import pymongo
import streamlit as st
import google.generativeai as genai
from pymongo.server_api import ServerApi
from llama_index.llms.gemini import Gemini
from google.api_core.exceptions import InvalidArgument
from llama_index.embeddings.gemini import GeminiEmbedding
from google.auth.exceptions import DefaultCredentialsError
from llama_index.core import VectorStoreIndex, Document, Settings, SummaryIndex

# Configure the API key
# genai.configure(api_key = st.secrets["gemini_api_key"])
st.set_page_config(page_title="PBoT")

@st.cache_resource(ttl=600)
def init_connection():
    username = urllib.parse.quote_plus(st.secrets["mongo"]["username"])
    password = urllib.parse.quote_plus(st.secrets["mongo"]["password"])
    return pymongo.MongoClient("mongodb+srv://%s:%s@cluster0.0hmur.mongodb.net/?retryWrites=true&w=majority&appName=Cluster0"%(username, password), server_api=ServerApi('1'))

client = init_connection()

db = client.pbot

# Authentication function for validating the API key
def validate_api_key(api_key):
    try:
        # Attempt to create a Gemini instance with the provided API key
        Gemini(api_key) # Assuming this initiates the client and checks the key
        return True
    
    except InvalidArgument or DefaultCredentialsError:
        return False

if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state["authenticated"]:
    st.markdown('<h2 class="main-title", align="center">Gemini API Authentication</h2>', unsafe_allow_html=True)
    st.markdown('<h3 class="sub-title" align="center">Please enter your API key to access the app</h3>', unsafe_allow_html=True)
    # Create a login box with custom styling
    with st.container():
        st.markdown('<div class="login-box">', unsafe_allow_html=True)
            
        # Text input for API key
        api_key = st.text_input("Gemini Api Key", type="password", help="Enter Your Gemini Api Key Here.")
            
        # Validate button
        login_button = st.button("Input Api Key")

        # Handle login button click
        if login_button:
            if validate_api_key(api_key):
                st.session_state['authenticated'] = True
                st.session_state['api_key'] = api_key
                st.rerun()
            else:
                st.write("Invalid Api Key. Please Try Again.")
    
    with st.container():
        st.markdown("<h3 align='center'>Get your API key from Google AI Studio.</h3>", unsafe_allow_html=True)
        
        st.markdown("""
            <p align='center'>To generate your API key, visit the <a href="https://ai.google.dev/gemini-api" target="_blank">
            Google AI Studio</a>.</p>
        """, unsafe_allow_html=True)


    st.divider()
    
    st.markdown("<p style='color:blue;text-align:center;'>Note: To use the model without api key leave the api field blank and continue.</p>", unsafe_allow_html=True)

else:

    fs = gridfs.GridFS(db)

    st.header("Chat With Your own Docs 💬 📚")

    with st.sidebar:
        with st.popover("Choose a model", use_container_width=True):
            model_list = [i.name.split("/")[-1] for i in genai.list_models()]
            # Create buttons for each model and store selected model in session_state
            for model_name in model_list:
                if st.button(model_name, use_container_width=True, key=f"model_{model_name}"):
                    st.session_state['selected_model'] = model_name

        st.divider()
        with st.popover("Uploaded Files", use_container_width=True):
            for file in fs._files.find():
                if st.button(file["filename"], key=f"file_{file['_id']}"):
                    st.session_state["file_id"] = file["_id"]

        st.divider()
        uploaded_file = st.file_uploader(label="Upload a file", accept_multiple_files=False, type=["docx","pdf","txt"])

        if uploaded_file:

            existing_file = fs.find_one({"filename": uploaded_file.name})
            if not existing_file:
                fs.put(uploaded_file.read(), filename=uploaded_file.name)
                st.rerun()

        st.divider()
        st.header("Get your API key from Google AI Studio.")
    
        # Link to the Gemini API page in Google AI Studio
        st.markdown("""
            <p>To generate your API key, visit the <a href="https://ai.google.dev/gemini-api" target="_blank">
            Google AI Studio</a>.</p>
        """, unsafe_allow_html=True)

        st.divider()
        logout_button = st.sidebar.button("Logout", use_container_width=True)
        if logout_button:
            st.session_state["authenticated"] = False
            st.session_state["api_key"] = None
            st.rerun()  # Rerun to force logout and show login page


    if "messages" not in st.session_state.keys():  # Initialize the chat message history
        st.session_state.messages = [
            {"role": "assistant", "content": "Ask me about yourself!"}
        ]


    # @st.cache_resource(show_spinner=False)
    def load_data():
        with st.spinner(text="Loading and indexing the files – hang on! This should take a few minutes."):

            if "selected_model" not in st.session_state.keys():
                st.session_state["selected_model"] = "gemini-1.5-flash"

            # Create the LLM (Gemini) instance directly
            Settings.llm = Gemini(api_key=st.secrets["gemini_api_key"], model=f"models/{st.session_state['selected_model']}", temperature=0.5)
            Settings.embed_model = GeminiEmbedding(model_name="models/embedding-001", api_key=st.secrets["gemini_api_key"])

            # Convert each file to a usable document for indexing                        
            if "file_id" not in st.session_state.keys():
                st.session_state["file_id"] =  list(fs._files.find({"filename":"readme.txt"}))[0]["_id"]

            # st.write(st.session_state["file_id"])

            # Retrieve the file content from fs.chunks using the file's _id
            file = fs.get(st.session_state["file_id"])

            # Convert binary data to text (assuming text files or similar)
            file_ext = file.filename.split(".")[-1].lower()

            if file_ext == "docx":
                doc = docx.Document(io.BytesIO(file.read()))
                file_content = "\n".join([para.text for para in doc.paragraphs])

            elif file_ext == "pdf":
                pdf = fitz.open(stream=io.BytesIO(file.read()), filetype="pdf")
                file_content = ""
                for page in range(pdf.page_count):
                    page = pdf.load_page(page)
                    file_content += page.get_text("text")
                    
            else:
                file_content = file.read().decode("utf-8", errors="ignore")            

            # Create a document with content and metadata (adjust fields as needed)
            text_list = [file_content]
            docs = [Document(text=t) for t in text_list]

            # Directly use the LLM when creating the VectorStoreIndex
            index = SummaryIndex.from_documents(docs, llm=Settings.llm, embed_model=Settings.embed_model)
            return index
    
    index = load_data()

    chat_engine = index.as_chat_engine(chat_mode="condense_question", verbose=True)

    if prompt := st.chat_input("Your question"): # Prompt for user input and save to chat history
        st.session_state.messages.append({"role": "user", "content": prompt})

    for message in st.session_state.messages: # Display the prior chat messages
        with st.chat_message(message["role"]):
            st.write(message["content"])

    # If last message is not from assistant, generate a new response
    if st.session_state.messages[-1]["role"] != "assistant":
        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                response = chat_engine.chat(prompt)
                st.write(response.response)
                message = {"role": "assistant", "content": response.response}
                st.session_state.messages.append(message) # Add response to message history